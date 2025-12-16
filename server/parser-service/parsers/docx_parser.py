from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from typing import Dict, Any, List
from .base_parser import BaseParser
import re

class DOCXParser(BaseParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        result = self.create_result_structure()
        
        try:
            doc = Document(file_path)
            
            # Расширенные метаданные
            result['metadata'] = {
                'type': 'docx',
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'keywords': doc.core_properties.keywords or '',
                'category': doc.core_properties.category or '',
                'comments': doc.core_properties.comments or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'last_modified_by': doc.core_properties.last_modified_by or '',
                'revision': doc.core_properties.revision,
            }
            
            full_text = []
            structure = []
            tables = []
            images = []
            lists = []
            
            # Проход по всем элементам документа в правильном порядке
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    para = self._find_paragraph(doc, element)
                    
                    if para and para.text.strip():
                        text = para.text.strip()
                        full_text.append(text)
                        
                        # Определение типа элемента
                        element_type, level = self._classify_paragraph(para)
                        
                        # Анализ форматирования текста
                        formatting = self._extract_formatting(para)
                        
                        structure.append({
                            'type': element_type,
                            'text': text,
                            'style': para.style.name,
                            'level': level,
                            'formatting': formatting,
                            'alignment': str(para.alignment) if para.alignment else 'LEFT',
                        })
                        
                        # Проверка на список
                        if self._is_list_item(para):
                            lists.append({
                                'text': text,
                                'level': level,
                                'numbered': self._is_numbered_list(para)
                            })
                
                elif isinstance(element, CT_Tbl):
                    table = self._find_table(doc, element)
                    
                    if table:
                        table_data = self._extract_table(table, len(tables))
                        tables.append(table_data)
            
            # Извлечение изображений
            images = self._extract_images(doc)
            
            result['content']['text'] = '\n'.join(full_text)
            result['content']['structure'] = structure
            result['content']['tables'] = tables
            result['content']['images'] = images
            result['content']['lists'] = lists
            
            # Статистика
            result['metadata']['paragraph_count'] = len([s for s in structure if s['type'] == 'paragraph'])
            result['metadata']['heading_count'] = len([s for s in structure if s['type'].startswith('heading')])
            result['metadata']['table_count'] = len(tables)
            result['metadata']['image_count'] = len(images)
            result['metadata']['list_count'] = len(lists)
            result['metadata']['word_count'] = len(' '.join(full_text).split())
            
        except Exception as e:
            result['metadata']['error'] = str(e)
        
        return result
    
    def _find_paragraph(self, doc: Document, element) -> Paragraph:
        """Найти объект параграфа по элементу"""
        for p in doc.paragraphs:
            if p._element == element:
                return p
        return None
    
    def _find_table(self, doc: Document, element) -> Table:
        """Найти объект таблицы по элементу"""
        for t in doc.tables:
            if t._element == element:
                return t
        return None
    
    def _classify_paragraph(self, para: Paragraph) -> tuple:
        """Классификация параграфа"""
        style = para.style.name.lower()
        
        if 'heading' in style:
            # Извлечь уровень заголовка
            level = 1
            if style.split()[-1].isdigit():
                level = int(style.split()[-1])
            return f'heading_{level}', level
        elif 'title' in style:
            return 'title', 0
        elif 'subtitle' in style:
            return 'subtitle', 0
        elif 'quote' in style or 'caption' in style:
            return 'quote', 0
        else:
            return 'paragraph', 0
    
    def _extract_formatting(self, para: Paragraph) -> Dict[str, Any]:
        """Извлечение форматирования текста"""
        formatting = {
            'bold': False,
            'italic': False,
            'underline': False,
            'font_sizes': [],
            'font_names': [],
            'colors': []
        }
        
        for run in para.runs:
            if run.bold:
                formatting['bold'] = True
            if run.italic:
                formatting['italic'] = True
            if run.underline:
                formatting['underline'] = True
            if run.font.size:
                size_pt = run.font.size.pt
                if size_pt not in formatting['font_sizes']:
                    formatting['font_sizes'].append(size_pt)
            if run.font.name:
                if run.font.name not in formatting['font_names']:
                    formatting['font_names'].append(run.font.name)
        
        return formatting
    
    def _is_list_item(self, para: Paragraph) -> bool:
        """Проверка, является ли параграф элементом списка"""
        text = para.text.strip()
        # Проверка маркеров списка
        return bool(re.match(r'^[\d\.\)]+\s+', text) or re.match(r'^[•\-\*◦▪▫]\\s+', text))
    
    def _is_numbered_list(self, para: Paragraph) -> bool:
        """Проверка, является ли список нумерованным"""
        text = para.text.strip()
        return bool(re.match(r'^[\d\.\)]+\s+', text))
    
    def _extract_table(self, table: Table, table_idx: int) -> Dict[str, Any]:
        """Улучшенное извлечение таблицы"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        
        # Определение заголовков (первая строка)
        headers = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []
        
        return {
            'table_index': table_idx,
            'headers': headers,
            'rows': data_rows,
            'row_count': len(table.rows),
            'col_count': len(table.columns),
            'all_rows': rows  # Включая заголовки
        }
    
    def _extract_images(self, doc: Document) -> List[Dict[str, Any]]:
        """Извлечение информации об изображениях"""
        images = []
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images.append({
                        'filename': rel.target_ref.split('/')[-1],
                        'content_type': rel.reltype
                    })
        except Exception as e:
            pass
        
        return images
